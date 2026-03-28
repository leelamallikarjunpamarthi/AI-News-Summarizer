"""
document_service.py
───────────────────
Handles file parsing for PDF, DOCX, and TXT files.
Returns clean, normalised text along with page-level metadata.
"""

from __future__ import annotations
import io
import re
from dataclasses import dataclass, field
import time
from typing import Optional
from loguru import logger

class OCRBackgroundRequired(Exception):
    """Raised when a document requires OCR and is too large for synchronous processing."""
    def __init__(self, message: str, total_pages: int):
        self.message = message
        self.total_pages = total_pages
        super().__init__(self.message)

try:
    import pdfplumber
    _PDFPLUMBER_AVAILABLE = True
except ImportError:
    _PDFPLUMBER_AVAILABLE = False

try:
    from pypdf import PdfReader
    _PYPDF_AVAILABLE = True
except ImportError:
    _PYPDF_AVAILABLE = False

try:
    import docx
    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False


@dataclass
class PageData:
    page_number: int
    text: str
    char_count: int = field(init=False)

    def __post_init__(self):
        self.char_count = len(self.text)


@dataclass
class ExtractedDocument:
    full_text: str
    pages: list[PageData]
    total_pages: int
    total_chars: int = field(init=False)

    def __post_init__(self):
        self.total_chars = len(self.full_text)


def _clean_text(text: str) -> str:
    """
    Normalise extracted text:
    - Collapse excessive whitespace
    - Remove form-feed characters
    - Fix broken hyphenation across lines
    """
    if not text:
        return ""
    text = text.replace("\f", "\n")
    text = re.sub(r"-\n(\w)", r"\1", text)          # Rejoin hyphenated words
    text = re.sub(r"[ \t]{2,}", " ", text)           # Collapse horizontal whitespace
    text = re.sub(r"\n{3,}", "\n\n", text)            # Max 2 consecutive newlines
    return text.strip()


def _extract_with_pypdf(file_bytes: bytes) -> ExtractedDocument:
    """Primary extraction strategy using PyPDF (much faster than pdfplumber)."""
    pages: list[PageData] = []
    reader = PdfReader(io.BytesIO(file_bytes))
    for i, page in enumerate(reader.pages, start=1):
        raw = page.extract_text() or ""
        cleaned = _clean_text(raw)
        pages.append(PageData(page_number=i, text=cleaned))

    full_text = "\n\n".join(p.text for p in pages if p.text)
    return ExtractedDocument(full_text=full_text, pages=pages, total_pages=len(pages))

def _extract_with_pdfplumber(file_bytes: bytes) -> ExtractedDocument:
    """Fallback extraction strategy using pdfplumber."""
    pages: list[PageData] = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for i, page in enumerate(pdf.pages, start=1):
            raw = page.extract_text() or ""
            cleaned = _clean_text(raw)
            pages.append(PageData(page_number=i, text=cleaned))

    full_text = "\n\n".join(p.text for p in pages if p.text)
    return ExtractedDocument(full_text=full_text, pages=pages, total_pages=len(pages))

def _extract_with_ocr(file_bytes: bytes, max_pages: int = 20) -> ExtractedDocument:
    """
    Final fallback: Convert PDF pages to images and use OCR (EasyOCR).
    Requires 'pymupdf' (fitz) and 'easyocr'.
    """
    try:
        import fitz  # PyMuPDF
        import easyocr
        import numpy as np
        from PIL import Image
    except ImportError as e:
        logger.error(f"OCR dependencies missing: {e}. Install with 'pip install pymupdf easyocr'.")
        raise ValueError("OCR extraction is not available due to missing dependencies.")

    logger.info("Initializing EasyOCR reader (this may take a few seconds on first run)...")
    # Initialize reader for English. 'gpu=False' if no GPU, but easyocr handles it usually.
    reader = easyocr.Reader(['en'])
    
    pages: list[PageData] = []
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    
    total_pages = len(doc)
    pages_to_process = min(total_pages, max_pages)
    
    logger.info(f"Performing OCR on {pages_to_process} of {total_pages} pages (Limit: {max_pages})...")
    
    if total_pages > max_pages:
        logger.warning(f"Document '{total_pages}' pages exceeds OCR limit of {max_pages}. Only the first {max_pages} pages will be processed.")

    for i in range(1, pages_to_process + 1):
        page = doc.load_page(i-1)
        # Render page to image (300 DPI for good OCR quality)
        pix = page.get_pixmap(matrix=fitz.Matrix(300/72, 300/72))
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
        
        # Convert PIL Image to numpy array for EasyOCR
        img_np = np.array(img)
        
        # Perform OCR
        results = reader.readtext(img_np, detail=0) # detail=0 returns just the text list
        raw_text = " ".join(results)
        cleaned = _clean_text(raw_text)
        
        pages.append(PageData(page_number=i, text=cleaned))
        logger.debug(f"Page {i} OCR complete. Extracted {len(cleaned)} chars.")

    doc.close()
    full_text = "\n\n".join(p.text for p in pages if p.text)
    
    if total_pages > max_pages:
        full_text += f"\n\n[NOTICE: Document truncated. Only the first {max_pages} pages were processed via OCR.]"
        
    return ExtractedDocument(full_text=full_text, pages=pages, total_pages=total_pages)

def _extract_docx(file_bytes: bytes) -> ExtractedDocument:
    """Extract text from a DOCX file."""
    if not _DOCX_AVAILABLE:
        raise ValueError("docx library is not installed. Run `pip install python-docx`")
        
    doc = docx.Document(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    
    # We treat the entire docx as single page for chunking purposes
    full_text = "\n".join(paragraphs)
    cleaned = _clean_text(full_text)
    
    pages = [PageData(page_number=1, text=cleaned)]
    return ExtractedDocument(full_text=cleaned, pages=pages, total_pages=1)


def _extract_txt(file_bytes: bytes) -> ExtractedDocument:
    """Extract text from a TXT file."""
    try:
        raw_text = file_bytes.decode('utf-8')
    except UnicodeDecodeError:
        try:
            raw_text = file_bytes.decode('windows-1252')
        except Exception:
            raw_text = file_bytes.decode('utf-8', errors='ignore')
            
    cleaned = _clean_text(raw_text)
    pages = [PageData(page_number=1, text=cleaned)]
    return ExtractedDocument(full_text=cleaned, pages=pages, total_pages=1)


def extract_text(file_bytes: bytes, filename: str, content_type: str, ocr_threshold: int = 5) -> ExtractedDocument:
    """
    Extract and clean text from an uploaded file based on its type.
    """
    start_time = time.perf_counter()
    filename_lower = filename.lower()
    
    logger.info(f"Starting text extraction for '{filename}' ({content_type})")

    try:
        # Text File handling
        if content_type == "text/plain" or filename_lower.endswith(".txt"):
            logger.debug(f"Extracting TXT file: {filename}")
            result = _extract_txt(file_bytes)
        
        # Docx File handling
        elif "wordprocessingml" in content_type or filename_lower.endswith(".docx") or filename_lower.endswith(".doc"):
            logger.debug(f"Extracting DOCX/DOC file: {filename}")
            result = _extract_docx(file_bytes)
            
        # PDF File handling (Default)
        else:
            if not (_PDFPLUMBER_AVAILABLE or _PYPDF_AVAILABLE):
                raise ValueError("No PDF parsing library available. Install pdfplumber or pypdf.")

            best_result: Optional[ExtractedDocument] = None

            # Try fast PyPDF first
            if _PYPDF_AVAILABLE:
                try:
                    pypdf_result = _extract_with_pypdf(file_bytes)
                    logger.debug(f"PyPDF extracted {pypdf_result.total_chars} chars.")
                    best_result = pypdf_result
                    
                    # If PyPDF did a great job (> 500 chars), we don't need backup
                    if pypdf_result.total_chars > 500:
                        result = pypdf_result
                        return result_with_log(result, start_time, filename)
                except Exception as e:
                    logger.warning(f"PyPDF extraction failed: {e}")

            # Try pdfplumber as backup or if PyPDF was mediocre
            if _PDFPLUMBER_AVAILABLE:
                try:
                    plumber_result = _extract_with_pdfplumber(file_bytes)
                    logger.debug(f"pdfplumber extracted {plumber_result.total_chars} chars.")
                    
                    if not best_result or plumber_result.total_chars > best_result.total_chars:
                        best_result = plumber_result
                except Exception as e:
                    logger.warning(f"pdfplumber extraction failed: {e}")

            if best_result:
                if best_result.total_chars > 0:
                    result = best_result
                else:
                    logger.info(f"Standard extraction returned 0 chars for '{filename}'.")
                    total_pages = best_result.total_pages if best_result else 0
                    
                    if total_pages > ocr_threshold:
                        logger.info(f"Document '{filename}' ({total_pages} pages) exceeds OCR threshold ({ocr_threshold}). Background processing required.")
                        raise OCRBackgroundRequired(
                            f"Document requires OCR and exceeds synchronous threshold of {ocr_threshold} pages.",
                            total_pages=total_pages
                        )
                        
                    logger.info(f"Attempting synchronous OCR fallback for '{filename}'...")
                    try:
                        # Use default max_pages=20 for sync OCR as well
                        result = _extract_with_ocr(file_bytes)
                    except Exception as ocr_err:
                        logger.error(f"OCR extraction failed for '{filename}': {ocr_err}")
                        raise ValueError("The PDF appears to be empty or contains only image-based/scanned content, and OCR extraction failed.")
            else:
                raise ValueError("PDF extraction failed with all available parsers.")
        
        return result_with_log(result, start_time, filename)

    except Exception as e:
        duration = time.perf_counter() - start_time
        logger.error(f"Extraction failed for '{filename}' after {duration:.2f}s: {e}")
        raise

def result_with_log(result: ExtractedDocument, start_time: float, filename: str) -> ExtractedDocument:
    duration = time.perf_counter() - start_time
    logger.info(f"Finished extraction for '{filename}' in {duration:.2f}s. Total chars: {result.total_chars}")
    return result


def build_chunk_metadata(
    document_id: str,
    filename: str,
    pages: list[PageData],
    chunk_index: int,
    chunk_text: str,
) -> dict:
    """
    Build metadata dict for a chunk, inferring the source page number.

    Args:
        document_id: Unique document ID.
        filename: Original uploaded filename.
        pages: List of PageData from extraction.
        chunk_index: Sequential index of this chunk.
        chunk_text: The chunk text (used for page inference).

    Returns:
        Metadata dict to store alongside the chunk embedding.
    """
    # Try to find which page this chunk text most likely originated from
    source_page = _infer_page(chunk_text, pages)
    return {
        "document_id": document_id,
        "filename": filename,
        "chunk_index": chunk_index,
        "page": source_page,
    }


def _infer_page(chunk_text: str, pages: list[PageData]) -> int:
    """
    Heuristic: find the page whose text contains the chunk snippet.
    Falls back to a scoring system if exact match isn't found.
    """
    snippet = chunk_text[:150].strip()
    if not snippet:
        return 1

    # 1. Try exact snippet match (fastest)
    for page in pages:
        if snippet in page.text:
            return page.page_number

    # 2. Fallback to character overlap (slower, but covers partial matches)
    best_page = 1
    best_score = 0
    # Sample characters to avoid quadratic complexity on huge pages
    probe = snippet[::5] 
    for page in pages:
        score = sum(1 for ch in probe if ch in page.text)
        if score > best_score:
            best_score = score
            best_page = page.page_number
    return best_page

